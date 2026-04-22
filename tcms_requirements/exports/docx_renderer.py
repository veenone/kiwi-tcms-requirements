"""DOCX report renderer (python-docx).

Two report scopes:
 - `requirement_list` — table of filtered requirements + per-row status/priority
 - `dashboard` — coverage KPIs + status/level/category breakdown tables

Output is written into a BytesIO buffer so the view can stream it with
`Content-Disposition: attachment`. Styling is intentionally plain — teams
customise via the post-generation track-changes review they'd do anyway.
"""
import io
from datetime import datetime, timezone

from tcms_requirements import __version__


def _now_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def _add_kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for key, value in rows:
        row = table.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)
    return table


def build_requirement_list_docx(queryset, *, title="Requirements report") -> bytes:
    from docx import Document  # noqa: WPS433

    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Generated: {_now_iso()} · kiwitcms-requirements v{__version__}")

    reqs = list(queryset)
    doc.add_paragraph(f"Total requirements: {len(reqs)}")

    if not reqs:
        doc.add_paragraph("No requirements match the selected filters.")
        return _dump(doc)

    _add_heading(doc, "Requirements", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text = "Identifier"
    header[1].text = "Title"
    header[2].text = "Level"
    header[3].text = "Status"
    header[4].text = "Priority"
    header[5].text = "Linked cases"

    for r in reqs:
        row = table.add_row().cells
        row[0].text = r.identifier
        row[1].text = r.title
        row[2].text = r.level.name if r.level_id else "—"
        row[3].text = r.get_status_display()
        row[4].text = r.get_priority_display()
        row[5].text = str(r.case_links.count())

    _add_heading(doc, "Detail", level=1)
    for r in reqs:
        _add_heading(doc, f"{r.identifier} — {r.title}", level=2)
        _add_kv_table(doc, [
            ("Level", r.level.name if r.level_id else "—"),
            ("Category", r.category.name if r.category_id else "—"),
            ("Status", r.get_status_display()),
            ("Priority", r.get_priority_display()),
            ("Verification", r.get_verification_method_display()),
            ("Source", r.source.name if r.source_id else "—"),
            ("Source section", r.source_section or "—"),
            ("Doc id / revision", f"{r.doc_id or '—'} {r.doc_revision or ''}".strip()),
            ("ASIL / DAL / IEC62304", " / ".join(filter(None, [r.asil, r.dal, r.iec62304_class])) or "—"),
            ("JIRA", r.jira_issue_key or "—"),
        ])
        if r.description:
            doc.add_paragraph(r.description)
        if r.rationale:
            _add_heading(doc, "Rationale", level=3)
            doc.add_paragraph(r.rationale)
        if r.case_links.exists():
            _add_heading(doc, "Linked test cases", level=3)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            header = table.rows[0].cells
            header[0].text = "TC ID"
            header[1].text = "Link type"
            header[2].text = "Suspect?"
            for link in r.case_links.all():
                cells = table.add_row().cells
                cells[0].text = f"TC-{link.case_id}"
                cells[1].text = link.get_link_type_display()
                cells[2].text = "⚠ suspect" if link.suspect else ""

    return _dump(doc)


def build_dashboard_docx(snapshot: dict, *, title="Requirements dashboard snapshot") -> bytes:
    from docx import Document  # noqa: WPS433

    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Generated: {_now_iso()} · kiwitcms-requirements v{__version__}")

    _add_heading(doc, "Coverage", level=1)
    coverage = snapshot.get("coverage", {})
    _add_kv_table(doc, [
        ("Total requirements", snapshot.get("total", 0)),
        ("Coverage", f"{coverage.get('percent', 0)} % ({coverage.get('linked', 0)} / {coverage.get('total', 0)})"),
        ("Orphan requirements", snapshot.get("orphan_requirements", 0)),
        ("Suspect links", snapshot.get("suspect_link_count", 0)),
    ])

    _add_heading(doc, "By status", level=1)
    _add_count_table(doc, snapshot.get("by_status", []), "status")
    _add_heading(doc, "By priority", level=1)
    _add_count_table(doc, snapshot.get("by_priority", []), "priority")
    _add_heading(doc, "By level", level=1)
    _add_count_table(doc, snapshot.get("by_level", []), "name", extra_key="code")
    _add_heading(doc, "By category", level=1)
    _add_count_table(doc, snapshot.get("by_category", []), "category")

    safety = snapshot.get("safety", {}) or {}
    if any(safety.values()):
        _add_heading(doc, "Safety / criticality distribution", level=1)
        for key, label in (("asil", "ASIL"), ("dal", "DAL"), ("iec62304_class", "IEC 62304 Class")):
            bucket = safety.get(key) or {}
            if not bucket:
                continue
            doc.add_paragraph(label, style="Intense Quote")
            table = doc.add_table(rows=1, cols=2)
            table.style = "Light Grid Accent 1"
            header = table.rows[0].cells
            header[0].text = label
            header[1].text = "Count"
            for class_name, count in sorted(bucket.items()):
                row = table.add_row().cells
                row[0].text = str(class_name)
                row[1].text = str(count)

    return _dump(doc)


def _add_count_table(doc, rows, label_key, extra_key=None):
    from docx import Document  # noqa: F401,WPS433 — keeps import close to the helper
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text = label_key.title()
    header[1].text = "Count"
    for row in rows or []:
        cells = table.add_row().cells
        if extra_key and row.get(extra_key):
            cells[0].text = f"{row.get(label_key, '—')} ({row[extra_key]})"
        else:
            cells[0].text = str(row.get(label_key, "—"))
        cells[1].text = str(row.get("count", 0))


def build_traceability_docx(rows, *, title="Requirements traceability report", diagram_png=None) -> bytes:
    """Traceability export: Sankey image (if diagram_png bytes supplied) + table.

    `rows` is the flattened row list from `traceability.report.flatten_traceability`.
    `diagram_png` is optional PNG bytes — skipped if None.
    """
    from docx import Document  # noqa: WPS433
    from docx.shared import Inches

    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Generated: {_now_iso()} · kiwitcms-requirements v{__version__}")

    if diagram_png:
        _add_heading(doc, "Traceability diagram", level=1)
        doc.add_picture(io.BytesIO(diagram_png), width=Inches(6.5))
        doc.add_paragraph(
            "Rendered from the browser view at the time of export. Blue = requirements, "
            "orange = test cases, green = test plans, red strokes = suspect links."
        )

    _add_heading(doc, "Traceability table", level=1)
    if not rows:
        doc.add_paragraph("No traceability rows match the current filters.")
        return _dump(doc)

    doc.add_paragraph(f"{len(rows)} row(s).")
    table = doc.add_table(rows=1, cols=7)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text = "Requirement"
    header[1].text = "Title"
    header[2].text = "Level"
    header[3].text = "Link"
    header[4].text = "Test case"
    header[5].text = "Test plan"
    header[6].text = "Suspect?"

    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row["req_identifier"]
        cells[1].text = row["req_title"]
        cells[2].text = row["level"] or "—"
        cells[3].text = row["link_type"] or "—"
        cells[4].text = f"TC-{row['case_id']}" if row["case_id"] else "—"
        cells[5].text = row["plan_name"] or "—"
        cells[6].text = "⚠ suspect" if row["suspect"] else ""

    return _dump(doc)


def _dump(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
