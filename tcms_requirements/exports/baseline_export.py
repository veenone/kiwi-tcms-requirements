"""Frozen-state exports for ProjectBaseline (DOCX / PDF / CSV).

These render the snapshot tables, NOT the live Requirement / RequirementTestCaseLink
rows. That preserves audit replay even after the source rows are edited or
deleted.
"""
import csv
import io
from datetime import datetime, timezone

from tcms_requirements import __version__


def _now_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")


def _baseline_metadata_rows(baseline) -> list:
    project = baseline.project
    return [
        ("Baseline name", baseline.name),
        ("Project", project.name),
        ("Project code", project.code or "—"),
        ("Product", str(project.product)),
        ("Kiwi Version", str(baseline.version) if baseline.version_id else "—"),
        ("Created at", baseline.created_at.strftime("%Y-%m-%d %H:%M UTC")),
        (
            "Created by",
            (baseline.created_by.get_full_name() or baseline.created_by.username)
            if baseline.created_by_id else "—",
        ),
        ("Notes", baseline.notes or "—"),
    ]


def build_baseline_csv(baseline, snapshots, link_snapshots) -> bytes:
    """One CSV with metadata as commented header, then frozen requirement rows."""
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow([f"# Baseline: {baseline.name}"])
    writer.writerow([f"# Project: {baseline.project.name}"])
    writer.writerow([f"# Created: {baseline.created_at.isoformat()}"])
    writer.writerow([f"# Plugin version: {__version__}"])
    writer.writerow([])
    writer.writerow([
        "identifier", "title", "status", "priority", "level_code",
        "asil", "sil", "iec62304_class", "dal", "linked_cases",
    ])

    case_count = {}
    for link in link_snapshots:
        case_count[link.requirement_identifier] = (
            case_count.get(link.requirement_identifier, 0) + 1
        )

    for snap in snapshots:
        writer.writerow([
            snap.identifier, snap.title, snap.status, snap.priority,
            snap.level_code, snap.asil, snap.sil, snap.iec62304_class, snap.dal,
            case_count.get(snap.identifier, 0),
        ])
    return buf.getvalue().encode("utf-8")


def _add_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (label, value) in enumerate(rows):
        c = table.rows[i].cells
        c[0].text = str(label)
        c[1].text = str(value) if value not in (None, "") else "—"


def build_baseline_docx(baseline, snapshots, link_snapshots) -> bytes:
    """Frozen requirement table + link table for one baseline."""
    from docx import Document  # noqa: WPS433

    doc = Document()
    doc.add_heading(f"Baseline: {baseline.name}", level=0)
    doc.add_paragraph(f"Generated: {_now_iso()} · kiwitcms-requirements v{__version__}")

    doc.add_heading("Metadata", level=1)
    _add_kv_table(doc, _baseline_metadata_rows(baseline))

    doc.add_heading("Frozen requirements", level=1)
    doc.add_paragraph(f"{len(snapshots)} requirement(s) at baseline time.")
    if snapshots:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        header = table.rows[0].cells
        for idx, label in enumerate([
            "Identifier", "Title", "Status", "Priority", "Level", "Safety",
        ]):
            header[idx].text = label
        for snap in snapshots:
            cells = table.add_row().cells
            cells[0].text = snap.identifier
            cells[1].text = snap.title
            cells[2].text = snap.status
            cells[3].text = snap.priority or "—"
            cells[4].text = snap.level_code or "—"
            safety_bits = [b for b in (
                snap.asil and f"ASIL {snap.asil}",
                snap.sil and f"SIL {snap.sil}",
                snap.iec62304_class and f"IEC62304 {snap.iec62304_class}",
                snap.dal and f"DAL {snap.dal}",
            ) if b]
            cells[5].text = ", ".join(safety_bits) or "—"

    doc.add_heading("Frozen test-case links", level=1)
    doc.add_paragraph(f"{len(link_snapshots)} link(s) at baseline time.")
    if link_snapshots:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        for idx, label in enumerate(["Requirement", "Case ID", "Link type", "Suspect"]):
            table.rows[0].cells[idx].text = label
        for link in link_snapshots:
            cells = table.add_row().cells
            cells[0].text = link.requirement_identifier
            cells[1].text = f"TC-{link.case_id}"
            cells[2].text = link.link_type
            cells[3].text = "Yes" if link.suspect else "No"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_baseline_pdf(baseline, snapshots, link_snapshots) -> bytes:
    """Frozen requirement table + link table for one baseline."""
    from reportlab.lib import colors  # noqa: WPS433
    from reportlab.lib.pagesizes import A4  # noqa: WPS433
    from reportlab.lib.styles import getSampleStyleSheet  # noqa: WPS433
    from reportlab.platypus import (  # noqa: WPS433
        Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
    )

    styles = getSampleStyleSheet()
    buf = io.BytesIO()
    pdf = SimpleDocTemplate(buf, pagesize=A4, title=f"Baseline {baseline.name}")
    story = [
        Paragraph(f"Baseline: {baseline.name}", styles["Title"]),
        Paragraph(
            f"Generated: {_now_iso()} · kiwitcms-requirements v{__version__}",
            styles["Italic"],
        ),
        Spacer(1, 12),
    ]

    story.append(Paragraph("Metadata", styles["Heading1"]))
    meta_data = [[k, str(v)] for k, v in _baseline_metadata_rows(baseline)]
    meta_table = Table(meta_data, colWidths=[120, 360])
    meta_table.setStyle(TableStyle([
        ("BOX", (0, 0), (-1, -1), 0.25, colors.grey),
        ("INNERGRID", (0, 0), (-1, -1), 0.1, colors.lightgrey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 12))

    story.append(Paragraph(
        f"Frozen requirements ({len(snapshots)})", styles["Heading1"]
    ))
    if snapshots:
        rows = [["Identifier", "Title", "Status", "Priority", "Level"]]
        for snap in snapshots:
            rows.append([
                snap.identifier,
                Paragraph(snap.title, styles["BodyText"]),
                snap.status,
                snap.priority or "—",
                snap.level_code or "—",
            ])
        req_table = Table(rows, colWidths=[80, 240, 60, 50, 50], repeatRows=1)
        req_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("BOX", (0, 0), (-1, -1), 0.25, colors.grey),
            ("INNERGRID", (0, 0), (-1, -1), 0.1, colors.lightgrey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(req_table)
    story.append(Spacer(1, 12))

    story.append(Paragraph(
        f"Frozen test-case links ({len(link_snapshots)})", styles["Heading1"]
    ))
    if link_snapshots:
        rows = [["Requirement", "Case", "Link type", "Suspect"]]
        for link in link_snapshots:
            rows.append([
                link.requirement_identifier,
                f"TC-{link.case_id}",
                link.link_type,
                "Yes" if link.suspect else "No",
            ])
        link_table = Table(rows, colWidths=[140, 80, 100, 60], repeatRows=1)
        link_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("BOX", (0, 0), (-1, -1), 0.25, colors.grey),
            ("INNERGRID", (0, 0), (-1, -1), 0.1, colors.lightgrey),
        ]))
        story.append(link_table)

    pdf.build(story)
    return buf.getvalue()
